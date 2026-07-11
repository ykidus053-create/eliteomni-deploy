"""Reliable multimodal gateway for EliteOmni.

Loads credentials from the deployment environment or repository-root .env,
uses provider-correct vision/OCR payloads, falls back from Mistral vision to
Groq vision, and provides cached document OCR.
"""
from __future__ import annotations

import base64
import binascii
import hashlib
import io
import mimetypes
import os
import re
import threading
import time
from collections import OrderedDict
from pathlib import Path
from typing import Any, Mapping

try:
    import requests
except ImportError as exc:  # pragma: no cover
    raise RuntimeError(
        "The requests package is required for multimodal calls"
    ) from exc


_REPO_ROOT = Path(__file__).resolve().parents[1]
_ENV_LOCK = threading.RLock()
_ENV_LOADED = False
_CACHE_LOCK = threading.RLock()
_OCR_CACHE: "OrderedDict[str, tuple[float, str]]" = OrderedDict()
_OCR_CACHE_MAX = 32
_LAST_STATUS: dict[str, Any] = {
    "vision_provider": None,
    "vision_error": None,
    "ocr_error": None,
}


class ProviderHTTPError(RuntimeError):
    def __init__(self, provider: str, status: int, body: str):
        self.provider = provider
        self.status = int(status)
        self.body = body[:800]
        super().__init__(f"{provider} HTTP {status}: {self.body}")


def _load_env_files() -> None:
    """Load dotenv-style files without overriding process environment."""
    global _ENV_LOADED
    with _ENV_LOCK:
        if _ENV_LOADED:
            return

        candidates = [
            _REPO_ROOT / ".env",
            Path.cwd() / ".env",
            _REPO_ROOT / "modules" / ".env",  # legacy location
        ]
        seen: set[Path] = set()

        for candidate in candidates:
            try:
                resolved = candidate.resolve()
            except OSError:
                resolved = candidate

            if resolved in seen or not candidate.is_file():
                continue
            seen.add(resolved)

            try:
                lines = candidate.read_text(encoding="utf-8").splitlines()
            except OSError:
                continue

            for raw in lines:
                line = raw.strip()
                if not line or line.startswith("#"):
                    continue
                if line.startswith("export "):
                    line = line[7:].lstrip()
                if "=" not in line:
                    continue

                key, value = line.split("=", 1)
                key = key.strip()
                value = value.strip()
                if (
                    len(value) >= 2
                    and value[0] == value[-1]
                    and value[0] in {'"', "'"}
                ):
                    value = value[1:-1]
                if key:
                    os.environ.setdefault(key, value)

        _ENV_LOADED = True


def _first_secret(*names: str) -> tuple[str, str | None]:
    _load_env_files()
    for name in names:
        value = os.environ.get(name, "").strip()
        if value:
            return value, name
    return "", None


def _mistral_key() -> tuple[str, str | None]:
    return _first_secret(
        "MISTRAL_OCR_API_KEY",
        "MISTRAL_VISION_API_KEY",
        "MISTRAL_API_KEY",
        "MISTRAL_KEY",
        "MISTRAL_TOKEN",
    )


def _groq_key() -> tuple[str, str | None]:
    return _first_secret(
        "GROQ_VISION_API_KEY",
        "GROQ_API_KEY",
    )


def _integer_env(
    name: str,
    default: int,
    minimum: int,
    maximum: int,
) -> int:
    try:
        value = int(os.environ.get(name, str(default)))
    except ValueError:
        value = default
    return max(minimum, min(value, maximum))


def _float_env(
    name: str,
    default: float,
    minimum: float,
    maximum: float,
) -> float:
    try:
        value = float(os.environ.get(name, str(default)))
    except ValueError:
        value = default
    return max(minimum, min(value, maximum))


def _strip_data_url(value: str) -> tuple[str, str | None]:
    if not isinstance(value, str):
        raise ValueError("base64 input must be a string")

    text = value.strip()
    mime = None
    match = re.match(
        r"^data:([^;,]+)(?:;[^,]*)?;base64,(.*)$",
        text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if match:
        mime = match.group(1).lower()
        text = match.group(2)

    text = re.sub(r"\s+", "", text)
    if not text:
        raise ValueError("base64 input is empty")

    text += "=" * (-len(text) % 4)
    try:
        decoded = base64.b64decode(text, validate=True)
    except (binascii.Error, ValueError) as exc:
        raise ValueError("invalid base64 input") from exc

    max_bytes = _integer_env(
        "ELITE_MULTIMODAL_MAX_BYTES",
        50 * 1024 * 1024,
        1024,
        200 * 1024 * 1024,
    )
    if len(decoded) > max_bytes:
        raise ValueError(
            f"file is too large ({len(decoded)} bytes; limit {max_bytes})"
        )

    return base64.b64encode(decoded).decode("ascii"), mime


def _mime_for(
    filename: str,
    supplied: str | None,
    encoded_b64: str,
) -> str:
    if supplied:
        return supplied

    name = (filename or "").lower()
    known = {
        ".pdf": "application/pdf",
        ".docx": (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        ".pptx": (
            "application/vnd.openxmlformats-officedocument."
            "presentationml.presentation"
        ),
        ".doc": "application/msword",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".avif": "image/avif",
    }
    for suffix, mime in known.items():
        if name.endswith(suffix):
            return mime

    guessed = mimetypes.guess_type(filename or "")[0]
    if guessed:
        return guessed

    prefix = base64.b64decode(encoded_b64[:40] + "===")
    if prefix.startswith(b"%PDF"):
        return "application/pdf"
    if prefix.startswith(b"\x89PNG"):
        return "image/png"
    if prefix.startswith(b"\xff\xd8\xff"):
        return "image/jpeg"
    if prefix.startswith(b"RIFF") and b"WEBP" in prefix:
        return "image/webp"

    return "application/octet-stream"


def _data_url(encoded: str, mime: str) -> str:
    return f"data:{mime};base64,{encoded}"


def _session() -> requests.Session:
    session = requests.Session()
    adapter = requests.adapters.HTTPAdapter(
        pool_connections=4,
        pool_maxsize=8,
        max_retries=0,
        pool_block=False,
    )
    session.mount("https://", adapter)
    session.headers.update(
        {
            "Accept": "application/json",
            "Accept-Encoding": "gzip, deflate",
            "Content-Type": "application/json",
            "User-Agent": "EliteOmni-Multimodal-V22/1.0",
        }
    )
    return session


def _request_json(
    *,
    provider: str,
    url: str,
    key: str,
    payload: Mapping[str, Any],
    timeout: float,
) -> dict[str, Any]:
    if not key:
        raise RuntimeError(f"{provider} API key is not configured")

    attempts = _integer_env(
        "ELITE_MULTIMODAL_ATTEMPTS",
        3,
        1,
        5,
    )
    last: Exception | None = None

    for attempt in range(attempts):
        try:
            with _session() as session:
                response = session.post(
                    url,
                    headers={"Authorization": f"Bearer {key}"},
                    json=dict(payload),
                    timeout=(10, timeout),
                )

            if response.status_code == 200:
                data = response.json()
                if not isinstance(data, dict):
                    raise RuntimeError(
                        f"{provider} returned non-object JSON"
                    )
                return data

            failure = ProviderHTTPError(
                provider,
                response.status_code,
                response.text[:1200],
            )
            if response.status_code in {
                400,
                401,
                403,
                404,
                422,
            }:
                raise failure
            if response.status_code not in {
                408,
                409,
                429,
                500,
                502,
                503,
                504,
            }:
                raise failure
            last = failure

        except ProviderHTTPError:
            raise
        except Exception as exc:
            last = exc

        if attempt + 1 < attempts:
            time.sleep(min(8.0, (2**attempt) + 0.25))

    raise RuntimeError(f"{provider} request failed: {last}")


def _extract_chat_text(data: Mapping[str, Any]) -> str:
    choices = data.get("choices") or []
    if not choices:
        raise RuntimeError("provider returned no choices")

    message = choices[0].get("message") or {}
    content = message.get("content", "")

    if isinstance(content, str):
        result = content.strip()
    elif isinstance(content, list):
        result = "\n".join(
            str(block.get("text", "")).strip()
            for block in content
            if (
                isinstance(block, Mapping)
                and block.get("type") == "text"
            )
        ).strip()
    else:
        result = str(content).strip()

    if not result:
        raise RuntimeError("provider returned empty vision content")
    return result


def _mistral_vision(
    encoded: str,
    mime: str,
    prompt: str,
    max_tokens: int,
) -> str:
    key, _ = _mistral_key()
    if not key:
        raise RuntimeError("Mistral vision key is not configured")

    model = (
        os.environ.get(
            "MISTRAL_VISION_MODEL",
            "mistral-small-latest",
        ).strip()
        or "mistral-small-latest"
    )
    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        # Mistral expects a string, not a nested object.
                        "image_url": _data_url(encoded, mime),
                    },
                ],
            }
        ],
        "max_tokens": max_tokens,
        "temperature": 0.1,
    }

    data = _request_json(
        provider="Mistral vision",
        url="https://api.mistral.ai/v1/chat/completions",
        key=key,
        payload=payload,
        timeout=_float_env(
            "ELITE_VISION_TIMEOUT_SECONDS",
            60,
            10,
            180,
        ),
    )
    return _extract_chat_text(data)


def _groq_vision(
    encoded: str,
    mime: str,
    prompt: str,
    max_tokens: int,
) -> str:
    key, _ = _groq_key()
    if not key:
        raise RuntimeError("Groq vision key is not configured")

    model = os.environ.get(
        "GROQ_VISION_MODEL",
        "meta-llama/llama-4-scout-17b-16e-instruct",
    ).strip()
    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        # Groq follows the OpenAI nested image_url shape.
                        "image_url": {
                            "url": _data_url(encoded, mime),
                        },
                    },
                ],
            }
        ],
        "max_tokens": max_tokens,
        "temperature": 0.1,
    }

    data = _request_json(
        provider="Groq vision",
        url="https://api.groq.com/openai/v1/chat/completions",
        key=key,
        payload=payload,
        timeout=_float_env(
            "ELITE_VISION_TIMEOUT_SECONDS",
            60,
            10,
            180,
        ),
    )
    return _extract_chat_text(data)


def vision_describe_v22(
    image_b64: str,
    prompt: str = "",
) -> str:
    """Describe an image with Mistral-first/Groq-fallback auth."""
    global _LAST_STATUS

    try:
        encoded, supplied_mime = _strip_data_url(image_b64)
        mime = _mime_for("image", supplied_mime, encoded)
    except Exception as exc:
        return f"[vision unavailable: {exc}]"

    question = (
        prompt
        or "Describe this image accurately and concisely."
    ).strip()
    instruction = (
        "Inspect the image directly. Answer the user's question first. "
        "Transcribe visible text exactly when relevant. Distinguish observed "
        "content from inference and do not invent hidden details.\n\n"
        f"User question: {question}"
    )
    max_tokens = _integer_env(
        "ELITE_VISION_MAX_TOKENS",
        900,
        64,
        4000,
    )

    preference = (
        os.environ.get(
            "ELITE_VISION_PROVIDER",
            "auto",
        ).strip().lower()
    )
    providers = {
        "mistral": _mistral_vision,
        "groq": _groq_vision,
    }
    order = (
        ["groq", "mistral"]
        if preference == "groq"
        else ["mistral", "groq"]
    )
    if preference in providers:
        order = [preference] + [
            name for name in order if name != preference
        ]

    errors: list[str] = []
    for provider in order:
        key = (
            _mistral_key()[0]
            if provider == "mistral"
            else _groq_key()[0]
        )
        if not key:
            errors.append(f"{provider}: key not configured")
            continue

        try:
            result = providers[provider](
                encoded,
                mime,
                instruction,
                max_tokens,
            )
            _LAST_STATUS = {
                **_LAST_STATUS,
                "vision_provider": provider,
                "vision_error": None,
            }
            return result
        except ProviderHTTPError as exc:
            errors.append(f"{provider}: HTTP {exc.status}")
            _LAST_STATUS = {
                **_LAST_STATUS,
                "vision_error": str(exc),
            }
            continue
        except Exception as exc:
            errors.append(
                f"{provider}: {type(exc).__name__}: {exc}"
            )
            _LAST_STATUS = {
                **_LAST_STATUS,
                "vision_error": str(exc),
            }

    return "[vision unavailable: " + "; ".join(errors) + "]"


def _ocr_cache_key(
    encoded: str,
    filename: str,
    model: str,
) -> str:
    digest = hashlib.sha256(encoded.encode("ascii")).hexdigest()
    return f"{model}:{filename.lower()}:{digest}"


def _cache_get(key: str) -> str | None:
    ttl = _integer_env(
        "ELITE_OCR_CACHE_SECONDS",
        3600,
        0,
        86400,
    )
    if ttl <= 0:
        return None

    with _CACHE_LOCK:
        item = _OCR_CACHE.get(key)
        if not item:
            return None

        created, value = item
        if time.time() - created > ttl:
            _OCR_CACHE.pop(key, None)
            return None

        _OCR_CACHE.move_to_end(key)
        return value


def _cache_set(key: str, value: str) -> None:
    with _CACHE_LOCK:
        _OCR_CACHE[key] = (time.time(), value)
        _OCR_CACHE.move_to_end(key)
        while len(_OCR_CACHE) > _OCR_CACHE_MAX:
            _OCR_CACHE.popitem(last=False)


def _render_ocr_response(data: Mapping[str, Any]) -> str:
    pages = data.get("pages") or []
    rendered: list[str] = []

    for position, raw_page in enumerate(pages, start=1):
        page = raw_page if isinstance(raw_page, Mapping) else {}
        page_index = page.get("index")
        number = (
            int(page_index) + 1
            if isinstance(page_index, int)
            else position
        )

        parts = [f"## Page {number}"]
        header = str(page.get("header") or "").strip()
        markdown = str(page.get("markdown") or "").strip()
        footer = str(page.get("footer") or "").strip()

        if header:
            parts.append(header)
        if markdown:
            parts.append(markdown)

        for table in page.get("tables") or []:
            if not isinstance(table, Mapping):
                continue
            table_text = str(
                table.get("markdown")
                or table.get("content")
                or table.get("html")
                or ""
            ).strip()
            table_id = str(table.get("id") or "").strip()
            if table_text and table_text not in markdown:
                title = (
                    f"### Table {table_id}"
                    if table_id
                    else "### Table"
                )
                parts.extend([title, table_text])

        if footer:
            parts.append(footer)

        rendered.append("\n\n".join(parts))

    result = "\n\n---\n\n".join(rendered).strip()
    return result or "[OCR returned no text]"


def ocr_document_v22(
    file_b64: str,
    filename: str = "document.pdf",
) -> str:
    """Extract ordered Markdown from PDF, Office, or image input."""
    global _LAST_STATUS

    try:
        encoded, supplied_mime = _strip_data_url(file_b64)
        mime = _mime_for(filename, supplied_mime, encoded)
    except Exception as exc:
        return f"[OCR unavailable: {exc}]"

    key, key_source = _mistral_key()
    if not key:
        return (
            "[OCR unavailable: configure MISTRAL_API_KEY in the "
            "deployment environment or repository-root .env]"
        )

    model = (
        os.environ.get(
            "MISTRAL_OCR_MODEL",
            "mistral-ocr-latest",
        ).strip()
        or "mistral-ocr-latest"
    )
    cache_key = _ocr_cache_key(
        encoded,
        filename,
        model,
    )
    cached = _cache_get(cache_key)
    if cached is not None:
        return cached

    is_image = mime.startswith("image/")
    document = (
        {
            "type": "image_url",
            "image_url": _data_url(encoded, mime),
        }
        if is_image
        else {
            "type": "document_url",
            "document_url": _data_url(encoded, mime),
        }
    )

    base_payload: dict[str, Any] = {
        "model": model,
        "document": document,
        "include_image_base64": False,
    }
    rich_payload = {
        **base_payload,
        "table_format": os.environ.get(
            "MISTRAL_OCR_TABLE_FORMAT",
            "markdown",
        ),
        "extract_header": True,
        "extract_footer": True,
    }
    timeout = _float_env(
        "ELITE_OCR_TIMEOUT_SECONDS",
        120,
        15,
        600,
    )

    try:
        try:
            data = _request_json(
                provider="Mistral OCR",
                url="https://api.mistral.ai/v1/ocr",
                key=key,
                payload=rich_payload,
                timeout=timeout,
            )
        except ProviderHTTPError as exc:
            # Some OCR deployments reject newer optional fields.
            if exc.status not in {400, 422}:
                raise
            data = _request_json(
                provider="Mistral OCR",
                url="https://api.mistral.ai/v1/ocr",
                key=key,
                payload=base_payload,
                timeout=timeout,
            )

        result = _render_ocr_response(data)
        _cache_set(cache_key, result)
        _LAST_STATUS = {
            **_LAST_STATUS,
            "ocr_error": None,
            "ocr_key_source": key_source,
            "ocr_pages": len(data.get("pages") or []),
        }
        return result

    except ProviderHTTPError as exc:
        _LAST_STATUS = {
            **_LAST_STATUS,
            "ocr_error": str(exc),
        }
        if exc.status in {401, 403}:
            return (
                "[OCR authentication failed: replace the Mistral key "
                f"configured through "
                f"{key_source or 'MISTRAL_API_KEY'}]"
            )
        return (
            f"[OCR failed: HTTP {exc.status}: "
            f"{exc.body[:300]}]"
        )
    except Exception as exc:
        _LAST_STATUS = {
            **_LAST_STATUS,
            "ocr_error": str(exc),
        }
        return f"[OCR failed: {type(exc).__name__}: {exc}]"


def mistral_ocr_v22(
    file_b64: str,
    filename: str = "document.pdf",
) -> str:
    return ocr_document_v22(file_b64, filename)


def _local_pdf_text(data: bytes) -> str:
    try:
        import pypdf
    except ImportError:
        return ""

    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
        text = "\n\n".join(
            page.extract_text() or ""
            for page in reader.pages
        )
        return text.strip()
    except Exception:
        return ""


def _local_docx_text(data: bytes) -> str:
    try:
        import docx
    except ImportError:
        return ""

    try:
        document = docx.Document(io.BytesIO(data))
        parts = [
            paragraph.text
            for paragraph in document.paragraphs
        ]
        for table in document.tables:
            for row in table.rows:
                parts.append(
                    " | ".join(
                        cell.text for cell in row.cells
                    )
                )
        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_uploaded_file_v22(
    filename: str,
    data: bytes,
) -> str:
    """Extract upload text, using OCR for scans and images."""
    name = filename or "upload"
    suffix = Path(name.lower()).suffix
    max_chars = _integer_env(
        "ELITE_DOCUMENT_MAX_CHARS",
        250000,
        1000,
        1000000,
    )

    if not isinstance(data, (bytes, bytearray)):
        return "[Document error: uploaded content must be bytes]"
    raw = bytes(data)

    if suffix in {
        ".txt",
        ".md",
        ".py",
        ".js",
        ".ts",
        ".tsx",
        ".jsx",
        ".html",
        ".css",
        ".json",
        ".csv",
        ".yaml",
        ".yml",
        ".toml",
        ".sql",
        ".xml",
    }:
        return raw.decode(
            "utf-8",
            errors="replace",
        )[:max_chars]

    mode = (
        os.environ.get(
            "ELITE_DOCUMENT_OCR_MODE",
            "auto",
        ).strip().lower()
    )

    if suffix == ".pdf":
        local = _local_pdf_text(raw)
        threshold = _integer_env(
            "ELITE_PDF_LOCAL_TEXT_MIN_CHARS",
            120,
            0,
            10000,
        )
        if mode != "always" and len(local) >= threshold:
            return local[:max_chars]

        encoded = base64.b64encode(raw).decode("ascii")
        result = ocr_document_v22(encoded, name)
        if result.startswith("[OCR ") and local:
            return local[:max_chars]
        return result[:max_chars]

    if suffix in {".docx", ".doc"}:
        local = (
            _local_docx_text(raw)
            if suffix == ".docx"
            else ""
        )
        if mode != "always" and local:
            return local[:max_chars]

        encoded = base64.b64encode(raw).decode("ascii")
        result = ocr_document_v22(encoded, name)
        if result.startswith("[OCR ") and local:
            return local[:max_chars]
        return result[:max_chars]

    if suffix in {".pptx", ".ppt"}:
        encoded = base64.b64encode(raw).decode("ascii")
        return ocr_document_v22(
            encoded,
            name,
        )[:max_chars]

    if suffix in {
        ".png",
        ".jpg",
        ".jpeg",
        ".webp",
        ".avif",
    }:
        encoded = base64.b64encode(raw).decode("ascii")
        result = ocr_document_v22(encoded, name)
        if result.startswith("[OCR "):
            return vision_describe_v22(
                encoded,
                (
                    "Transcribe all visible text and describe "
                    "the document layout."
                ),
            )[:max_chars]
        return result[:max_chars]

    return (
        f"[Unsupported file type: "
        f"{suffix or 'unknown'}]"
    )


def multimodal_status() -> dict[str, Any]:
    mistral_key, mistral_source = _mistral_key()
    groq_key, groq_source = _groq_key()

    return {
        "mistral_configured": bool(mistral_key),
        "mistral_key_source": mistral_source,
        "groq_configured": bool(groq_key),
        "groq_key_source": groq_source,
        "vision_preference": os.environ.get(
            "ELITE_VISION_PROVIDER",
            "auto",
        ),
        "mistral_vision_model": os.environ.get(
            "MISTRAL_VISION_MODEL",
            "mistral-small-latest",
        ),
        "groq_vision_model": os.environ.get(
            "GROQ_VISION_MODEL",
            "meta-llama/llama-4-scout-17b-16e-instruct",
        ),
        "mistral_ocr_model": os.environ.get(
            "MISTRAL_OCR_MODEL",
            "mistral-ocr-latest",
        ),
        "document_ocr_mode": os.environ.get(
            "ELITE_DOCUMENT_OCR_MODE",
            "auto",
        ),
        "last_status": dict(_LAST_STATUS),
    }
